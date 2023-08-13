import csv
import matplotlib.pyplot as plt
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.ticker as ticker
from datetime import datetime
from docx.shared import Inches
import os


def string_to_float(s):
    """Convert a string with optional 'k', 'm', 'b' postfix to numeric values (thousand, million, billion)"""
    s = s.strip().lower()
    if s[-1] in ['k', 'm', 'b']:
        num, magnitude = s[:-1], s[-1]
        if magnitude == 'k':
            return float(num) * 1_000
        elif magnitude == 'm':
            return float(num) * 1_000_000
        elif magnitude == 'b':
            return float(num) * 1_000_000_000
    else:
        return float(s)

def format_number(number):
    """Convert a number into a string with 'k', 'm', 'b' postfix for thousands, millions, billions"""
    if number >= 1_000_000_000:
        return f"{number / 1_000_000_000:.0f}b"
    elif number >= 1_000_000:
        return f"{number / 1_000_000:.0f}m"
    elif number >= 1_000:
        return f"{number / 1_000:.0f}k"
    else:
        return f"{number:.0f}"

def percentage_change(old, new):
    old = string_to_float(old.replace(',', ''))
    new = string_to_float(new.replace(',', ''))
    if old == 0:
        return "0%"
    change = ((new - old) / old) * 100
    return f"{change:.0f}%"

def get_total_additional_engagements(initial_key, ending_key, indices, all_data):
    total_initial = sum(string_to_float(row[indices[initial_key]]) if row[indices[initial_key]].strip() else 0 for row in all_data)
    total_ending = sum(string_to_float(row[indices[ending_key]]) if row[indices[ending_key]].strip() else 0 for row in all_data)
    return total_ending - total_initial

def add_summary_table(document, data, title):
    table = document.add_table(rows=2, cols=len(data) + 1)
    table.style = 'Table Grid'
    table.cell(0, 0).text = title
    for i, (key, value) in enumerate(data.items(), start=1):
        table.cell(0, i).text = key
        table.cell(1, i).text = format_number(value)
    document.add_paragraph("\n")

def add_bar_graph(document, data, title, xlabel, ylabel, color):
    # Plot the bar graph using matplotlib
    fig, ax = plt.subplots()
    ax.bar(data.keys(), data.values(), color=color)
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)

    # Format y-axis labels
    def custom_formatter(x, _):
        return format_number(x)

    ax.yaxis.set_major_formatter(ticker.FuncFormatter(custom_formatter))

    # Save the bar graph to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    plt.savefig(temp_file.name, dpi=300, bbox_inches='tight')
    temp_file.close()

    # Add the bar graph as an image to the document
    document.add_picture(temp_file.name, width=document.sections[-1].page_width - document.sections[-1].left_margin - document.sections[-1].right_margin)
    plt.close()

    # Delete the temporary file
    temp_file.delete = True

def tables_and_comments(document, indices, row):
    # First table
    stats_table = document.add_table(rows=2, cols=6)
    stats_table.style = 'Table Grid'

    labels = ["", "Views", "Retweets", "Quotes", "Likes", "Bookmarks"]

    for i, label in enumerate(labels):
        stats_table.cell(0, i).text = label

    stats_table.cell(1, 0).text = "Initial"

    for i, key in enumerate(["Views", "Retweets", "Quotes", "Likes", "Bookmarks"]):
        stats_table.cell(1, i + 1).text = format_number(float(row[indices[key]]))

    document.add_paragraph("\n")

    # Second table
    ending_table = document.add_table(rows=3, cols=6)
    ending_table.style = 'Table Grid'

    for i, label in enumerate(labels):
        ending_table.cell(0, i).text = label

    ending_table.cell(1, 0).text = "Ending"
    ending_table.cell(2, 0).text = "% Change"

    for i, key in enumerate(["Ending Views", "Ending Retweets", "Ending Quotes", "Ending Likes", "Ending Bookmarks"]):
        ending_table.cell(1, i + 1).text = format_number(float(row[indices[key]]))
        ending_table.cell(2, i + 1).text = percentage_change(stats_table.cell(1, i + 1).text, ending_table.cell(1, i + 1).text)

    if include_comments.lower() == 'yes':
        comments = row[indices["Comments"]].split('\n')
        document.add_heading("Comments:", level=2)
        comments_list = document.add_paragraph()
        for i, comment in enumerate(comments, start=1):
            comments_list.add_run(f"{i}. {comment.strip()}\n")

    document.add_page_break()

input_file = input("Enter the CSV file path: ")
output_file = input("Enter the output DOCX file path: ")
include_comments = input("Do you wish to include comments in the report? [yes/no]: ")

# Remove ".docx", capitalize each word for report title
report_title = os.path.splitext(output_file)[0].replace('_', ' ').title()
date = datetime.today().strftime('%Y-%m-%d')  # Get today's date
  
document = Document()

# Add beego logo and center it
document.add_picture('beego.png', width=Inches(2.0))
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add report title and center it
title_paragraph = document.add_paragraph()
title_run = title_paragraph.add_run(report_title)
title_run.font.name = 'Helvetica Neue'
title_run.font.size = Pt(24)
title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add date and center it
date_paragraph = document.add_paragraph()
date_run = date_paragraph.add_run(date)
date_run.font.size = Pt(12)
date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.add_page_break()

# Set font for whole document
style = document.styles['Normal']
font = style.font
font.name = 'Helvetica Neue'
font.size = Pt(12)

# Set font for all heading styles used
for heading_level in range(1, 10):
    style = document.styles[f'Heading {heading_level}']
    font = style.font
    font.name = 'Helvetica Neue'
    font.size = Pt(12)

document.add_heading("Summary", level=1)

# Header and Footer
section = document.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_text = header_paragraph.add_run('LBC Swarm Campaign Report')
header_text.font.name = 'Helvetica Neue'
header_text.font.size = Pt(12)
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_text = footer_paragraph.add_run('Private and Confidential')
footer_text.font.name = 'Helvetica Neue'
footer_text.font.size = Pt(12)
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

with open(input_file, newline='') as csvfile:
    csv_data = csv.reader(csvfile)
    header = next(csv_data)

    indices = {name: header.index(name) for name in ["Swarm Number", "Swarm URL", "Views", "Retweets", "Quotes", "Likes", "Bookmarks", "Ending Views", "Ending Retweets", "Ending Quotes", "Ending Likes", "Ending Bookmarks", "Comments", "Swarm Week"]}

    all_data = list(csv_data)
    weeks = sorted(list(set(row[indices["Swarm Week"]] for row in all_data)))
    print(f"Available weeks: {', '.join(weeks)}")

    selected_week = input(f"Choose a week to create the report for, or enter 'all' for all weeks: ")

    if selected_week.lower() == 'all':
        report_data = all_data
    else:
        report_data = [row for row in all_data if row[indices["Swarm Week"]] == selected_week]

    total_comments = sum(1 for row in report_data for _ in row[indices['Comments']].split('\n') if row[indices['Comments']].strip())
    document.add_paragraph(f"Total Comments for Week {selected_week}: {format_number(total_comments)}")
    document.add_paragraph(f"Total Swarms for Week {selected_week}: {len(report_data)}")

    total_additional_engagements = {
        'Views': get_total_additional_engagements("Views", "Ending Views", indices, report_data),
        'Retweets': get_total_additional_engagements("Retweets", "Ending Retweets", indices, report_data),
        'Quotes': get_total_additional_engagements("Quotes", "Ending Quotes", indices, report_data),
        'Likes': get_total_additional_engagements("Likes", "Ending Likes", indices, report_data),
        'Bookmarks': get_total_additional_engagements("Bookmarks", "Ending Bookmarks", indices, report_data)
    }
    add_summary_table(document, total_additional_engagements, f"Total Additional Engagements for Week {selected_week}")

    average_additional_engagements = {key: value / len(report_data) for key, value in total_additional_engagements.items()}
    add_summary_table(document, average_additional_engagements, f"Total Average Engagements for Week {selected_week}")

    # Add bar graphs
    total_swarms_per_week = {week: sum(1 for row in all_data if row[indices["Swarm Week"]] == week) for week in weeks}
    add_bar_graph(document, total_swarms_per_week, "Total Swarms per Week", "Week", "Number of Swarms", color='darkorange')

    # Total number of Additional Engagements each week
    total_additional_engagements_per_week = {
        week: {key: get_total_additional_engagements(key, f"Ending {key}", indices, [row for row in report_data if row[indices["Swarm Week"]] == week])
               for key in ["Views", "Retweets", "Quotes", "Likes", "Bookmarks"]}
        for week in weeks
    }
    for eng_type in ["Views", "Retweets", "Quotes", "Likes", "Bookmarks"]:
        add_bar_graph(document, {week: data[eng_type] for week, data in total_additional_engagements_per_week.items()}, f"Total Additional {eng_type} per Week", "Week", f"Total Additional {eng_type}", color='darkorange')

    # Graph comparing the before and after engagements
    initial_total_engagements_per_week = {
        week: sum(sum(float(row[indices[key]]) for key in ["Views", "Retweets", "Quotes", "Likes", "Bookmarks"]) for row in all_data if row[indices["Swarm Week"]] == week)
        for week in weeks
    }
    ending_total_engagements_per_week = {
        week: sum(sum(float(row[indices[f"Ending {key}"]]) for key in ["Views", "Retweets", "Quotes", "Likes", "Bookmarks"]) for row in all_data if row[indices["Swarm Week"]] == week)
        for week in weeks
    }
    add_bar_graph(document, initial_total_engagements_per_week, "Initial Total Engagements per Week", "Week", "Initial Engagements", color='darkorange')
    add_bar_graph(document, ending_total_engagements_per_week, "Ending Total Engagements per Week", "Week", "Ending Engagements", color='darkorange')

    # Total number of comments per week
    total_comments_per_week = {
        week: sum(1 for row in all_data if row[indices["Swarm Week"]] == week for _ in row[indices['Comments']].split('\n') if row[indices['Comments']].strip())
        for week in weeks
    }
    add_bar_graph(document, total_comments_per_week, "Total Comments per Week", "Week", "Number of Comments", color='darkorange')

    for row in report_data:
        document.add_heading(f"Swarm Number: {row[indices['Swarm Number']]}", level=1)
        document.add_paragraph(f"Swarm URL: {row[indices['Swarm URL']]}")

        tables_and_comments(document, indices, row)

document.save(output_file)

print("Conversion complete! DOCX file saved as", output_file)